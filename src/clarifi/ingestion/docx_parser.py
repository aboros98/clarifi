from pathlib import Path

from docx import Document

from clarifi.ingestion.parser_factory import ParseResult


class DocxParser:
    async def parse(self, file_path: Path) -> ParseResult:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        return ParseResult(
            text="\n".join(paragraphs),
            page_count=None,
            ocr_applied=False,
        )
